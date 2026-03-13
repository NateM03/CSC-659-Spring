"""
Convert HW2_Report_Template.md to Word document
"""

try:
    import pypandoc
    # Convert markdown to docx
    output = pypandoc.convert_file(
        'HW2_Report_Template.md',
        'docx',
        outputfile='CSC 659-859 Spring 2026 HW 2 Moreno.docx',
        extra_args=['--standalone']
    )
    print("✅ Successfully converted to Word!")
    print("File saved as: CSC 659-859 Spring 2026 HW 2 Moreno.docx")
except ImportError:
    print("pypandoc not installed. Trying alternative method...")
    try:
        from docx import Document
        import re
        
        # Read markdown file
        with open('HW2_Report_Template.md', 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Create Word document
        doc = Document()
        
        # Simple markdown to docx conversion
        lines = content.split('\n')
        for line in lines:
            # Headers
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                doc.add_heading(line[5:], level=4)
            # Bold
            elif line.startswith('**') and line.endswith('**'):
                p = doc.add_paragraph()
                p.add_run(line[2:-2]).bold = True
            # Code blocks (skip for now, will be screenshots)
            elif line.startswith('```'):
                continue
            # Images
            elif line.startswith('!['):
                # Extract image path
                match = re.search(r'\(([^)]+)\)', line)
                if match:
                    img_path = match.group(1)
                    try:
                        doc.add_picture(img_path)
                    except:
                        p = doc.add_paragraph(f"[Image: {img_path}]")
            # Empty lines
            elif line.strip() == '':
                doc.add_paragraph()
            # Regular text
            else:
                # Handle inline formatting
                text = line
                # Bold
                text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
                p = doc.add_paragraph(text)
        
        doc.save('CSC 659-859 Spring 2026 HW 2 Moreno.docx')
        print("✅ Successfully converted to Word using python-docx!")
        print("File saved as: CSC 659-859 Spring 2026 HW 2 Moreno.docx")
        print("Note: Code blocks and complex formatting may need manual adjustment.")
    except ImportError:
        print("Neither pypandoc nor python-docx available.")
        print("\nAlternative options:")
        print("1. Install pandoc: https://pandoc.org/installing.html")
        print("2. Use online converter: https://cloudconvert.com/md-to-docx")
        print("3. Open in Word: Word can open .md files directly")
        print("4. Use VS Code: Install 'Markdown PDF' extension")

if __name__ == "__main__":
    print("Converting HW2_Report_Template.md to Word...")
    print("=" * 60)

